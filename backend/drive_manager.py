"""
יוצר תיקיות מאורגנות ב-Google Drive (G:\) עבור כל מכרז רלוונטי.
"""
import os
import re
import json
import shutil
from datetime import datetime

DRIVE_ROOT = r"G:\האחסון שלי\מכרזים"
COMPANY_DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "company_docs")
DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "documents")


def sanitize_folder_name(name: str) -> str:
    """מנקה שם לשימוש כשם תיקייה."""
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name[:60]


def get_submission_date_str(date_str: str) -> str:
    """ממיר תאריך הגשה לפורמט קצר עם נקודות (בטוח לשם תיקייה)."""
    if not date_str:
        return ""
    # נסה פורמטים שונים
    for fmt in ["%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d",
                "%d/%m/%Y %H:%M", "%d/%m/%Y"]:
        try:
            dt = datetime.strptime(date_str.strip()[:len(fmt)], fmt)
            return dt.strftime("%d.%m.%Y")
        except ValueError:
            continue
    # החזר בכל מקרה עם נקודות במקום סלאשים
    return date_str[:10].replace("/", ".").replace("-", ".")


def create_tender_folder(tender: dict, analysis: dict) -> str | None:
    """
    יוצר תיקייה ב-Google Drive עבור מכרז שעומד בתנאי סף.
    מחזיר את הנתיב לתיקייה.
    """
    if not os.path.exists(DRIVE_ROOT.split("\\")[0] + "\\"):
        print("Google Drive לא מחובר (G:\\)")
        return None

    # צור תיקיית שורש אם לא קיימת
    os.makedirs(DRIVE_ROOT, exist_ok=True)

    # שם התיקייה: שם מכרז + תאריך הגשה
    title = tender.get("title", "מכרז")
    date = get_submission_date_str(
        analysis.get("submission_deadline") or tender.get("submission_date", "")
    )
    publisher = tender.get("publisher", "")

    folder_name = sanitize_folder_name(f"{title} — {publisher}")
    if date:
        folder_name = f"[{date}] {folder_name}"

    folder_path = os.path.join(DRIVE_ROOT, folder_name)
    os.makedirs(folder_path, exist_ok=True)

    # יצירת קבצי המכרז בתיקייה
    _create_required_docs_file(folder_path, tender, analysis)
    _create_questions_file(folder_path, tender, analysis)
    _copy_company_docs(folder_path, analysis)
    _copy_tender_documents(folder_path, tender)

    print(f"Drive folder created: {folder_path}")
    return folder_path


def _create_required_docs_file(folder_path: str, tender: dict, analysis: dict):
    """יוצר קובץ Word עם רשימת מסמכים נדרשים."""
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(f"מסמכים נדרשים — {tender.get('title', '')}", 0)
        doc.add_paragraph(f"מפרסם: {tender.get('publisher', '')}")
        doc.add_paragraph(f"מועד הגשה: {analysis.get('submission_deadline', tender.get('submission_date', ''))}")
        doc.add_paragraph(f"עמידה בתנאי סף: {'✓ כן' if analysis.get('eligible') == 1 else '✗ לא'}")
        doc.add_paragraph("")

        doc.add_heading("תנאי סף", 1)
        for cond in analysis.get("threshold_conditions", []):
            status = "✓" if cond.get("met") else "✗"
            doc.add_paragraph(f"{status} {cond.get('condition', '')} — {cond.get('notes', '')}")

        doc.add_heading("מסמכים נדרשים להגשה", 1)
        for i, req in enumerate(analysis.get("required_documents", []), 1):
            doc.add_paragraph(f"{i}. {req}", style="List Number")

        doc.add_heading("סיכום", 1)
        doc.add_paragraph(analysis.get("ai_summary", ""))

        path = os.path.join(folder_path, "📋 מסמכים נדרשים.docx")
        doc.save(path)
    except Exception as e:
        print(f"Error creating docs file: {e}")


def _create_questions_file(folder_path: str, tender: dict, analysis: dict):
    """יוצר קובץ עם שאלות למזמין."""
    questions = analysis.get("questions_to_client", [])
    if not questions:
        return
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(f"שאלות למזמין — {tender.get('title', '')}", 0)
        doc.add_paragraph(f"מפרסם: {tender.get('publisher', '')}")
        doc.add_paragraph("")
        doc.add_heading("שאלות להגשה", 1)
        for i, q in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {q}", style="List Number")
        path = os.path.join(folder_path, "❓ שאלות למזמין.docx")
        doc.save(path)
    except Exception as e:
        print(f"Error creating questions file: {e}")


def _copy_company_docs(folder_path: str, analysis: dict):
    """מעתיק מסמכי חברה רלוונטיים לתיקייה."""
    if not os.path.exists(COMPANY_DOCS_DIR):
        return

    # תיקיית מסמכי חברה בתוך תיקיית המכרז
    company_subfolder = os.path.join(folder_path, "מסמכי חברה")
    os.makedirs(company_subfolder, exist_ok=True)

    for filename in os.listdir(COMPANY_DOCS_DIR):
        src = os.path.join(COMPANY_DOCS_DIR, filename)
        dst = os.path.join(company_subfolder, filename)
        if os.path.isfile(src) and not os.path.exists(dst):
            shutil.copy2(src, dst)
            print(f"  Copied: {filename}")


def _copy_tender_documents(folder_path: str, tender: dict):
    """מעתיק מסמכי מכרז שהורדו לתיקייה."""
    tender_id = tender.get("tender_id", "")
    if not tender_id:
        return

    safe_tid = re.sub(r'[^\w\-_]', '_', str(tender_id))
    source_dir = os.path.join(DOCS_DIR, safe_tid)
    if not os.path.exists(source_dir):
        return

    tender_docs_folder = os.path.join(folder_path, "מסמכי מכרז")
    os.makedirs(tender_docs_folder, exist_ok=True)

    for filename in os.listdir(source_dir):
        src = os.path.join(source_dir, filename)
        dst = os.path.join(tender_docs_folder, filename)
        if os.path.isfile(src) and not os.path.exists(dst):
            shutil.copy2(src, dst)


def setup_company_docs_folder():
    """יוצר תיקיית מסמכי חברה ומכניס הסברים."""
    os.makedirs(COMPANY_DOCS_DIR, exist_ok=True)
    readme_path = os.path.join(COMPANY_DOCS_DIR, "הוראות.txt")
    if not os.path.exists(readme_path):
        with open(readme_path, "w", encoding="utf-8") as f:
            f.write("""שים כאן את מסמכי החברה שיועתקו אוטומטית לכל תיקיית מכרז:

• פרופיל חברה.docx
• CV עקיבא רוכסר.docx
• אישור עוסק מורשה.pdf
• אישור ניהול פנקסי חשבונות.pdf
• רשימת פרויקטים.docx
• ממליצים.docx

המערכת תעתיק אותם אוטומטית לכל מכרז שעומד בתנאי הסף.
""")
    return COMPANY_DOCS_DIR
